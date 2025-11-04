'''
Dateikonverter.py
Author: ???
Description:
For now it only converts a file named zuKonvertierendeDatei.txt
into konvertierteDatei.pdf. With some more work, it could become a usefull file converter.
'''

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
#Nicht vergessen, die benötigten Pakete zu installieren:
#pip install python-docx reportlab

'''
Class Dateikonverter
Takes an input_file and tries to convert it into another format.
Methods:
_read_text:
  Reads the input file and returns its content as string.
to_markdown:
  Parameters: output_file (The name under which the converted file is saved)
  Converts the provided input_file to the markdown format.
to_word:
  Parameters: output_file (The name under which the converted file is saved)
  Converts the provided input_file to the docx format.
to_pdf:
  Parameters: output_file (The name under which the converted file is saved)
  Converts the provided input_file to the pdf format.
'''

class Dateikonverter:
    '''
    Constructor
    Parameters: input_file (The File you want to convert)
    '''
    def __init__(self, input_file):
        self.input_file = input_file

    def _read_text(self):
        '''
        Method: _read_text
        Reads the input_file and returns its content as string.
        '''
        with open(self.input_file, "r", encoding="utf-8") as f: # Opens the input_file for reading.
            return f.read() # Returns the content of the read file as string.

    def to_markdown(self, output_file):
        '''
        Method: to_markdown
        Parameters: output_file (The name under which the converted file is saved)
        Converts the provided input_file to the markdown format.
        '''
        text = self._read_text() # Read input_file and save as string in variable text.
        with open(output_file, "w", encoding="utf-8") as f: # Opens the output file for writing.
            f.write(f"# Converted File\n\n{text}") # Writes the content of the input_file into the new file.

    def to_word(self, output_file):
        '''
        Method: to_word
        Parameters: output_file (The name under which the converted file is saved)
        Converts the provided input_file to the docx format.
        '''
        text = self._read_text() # Read input_file and save as string in variable text.
        doc = Document() # Create instance from Document called doc.
        doc.add_heading("Converted File", level=1) # Probably adds Header for the the docx. Not sure. 
        doc.add_paragraph(text) # Adds the string which was read from input_file to the doc.
        doc.save(output_file) # Saves the Document under the filename provided by Parameter output_file

    def to_pdf(self, output_file):
        '''
        Method: to_pdf
        Parameters: output_file (The name under which the converted file is saved)
        Converts the provided input_file to the pdf format.
        '''
        text = self._read_text() # Read input_file and save as string in variable text.
        c = canvas.Canvas(output_file, pagesize=A4) # Creates a canvas of size DIN-A4
        width, height = A4 # Deconstructs A4 into width and height
        y = height - 50 # Space at the top of the Document?!
        '''
        Loop to write the rows of text onto the canvas
        '''
        for line in text.splitlines():
            c.drawString(50, y, line) #writes row with a width of 50?! at height y
            y -= 15 # space between the rows.
            if y < 50: # if space at the bottom of site is less then 50, start new page I assume.
                c.showPage()
                y = height - 50 # reset the height for the next page
        c.save() # save the canvas as pdf with the name provided by output file.

'''
Create an instance from Dateikonverter called test_konvertierer and
pass the file zuKonvertierendeDatei.txt as argument.
'''
test_konvertierer = Dateikonverter("zuKonvertierendeDatei.txt")

'''
Call the method test_konvertierer.to_pdf and pass the string 'konvertierteDatei'
Converts zuKonvertierendeDatei.txt to konvertierteDatei.pdf
'''
test_konvertierer.to_pdf("konvertierteDatei.pdf")
